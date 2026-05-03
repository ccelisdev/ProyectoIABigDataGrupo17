import os
import re
import logging
import pandas as pd
from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from src.embeddings import MotorEmbeddings
from src.database import GestorVectorial
from langchain_community.vectorstores import Qdrant

# Configuración de Logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

class IngestorDocumentos:
    def __init__(self):
        self.embeddings = MotorEmbeddings().obtener_modelo()
        self.gestor_db = GestorVectorial()
        self.cliente_qdrant = self.gestor_db.obtener_cliente()
        self.nombre_coleccion = self.gestor_db.nombre_coleccion

    def clean_text(self, text: str) -> str:
        if not text: return ""
        text = text.replace("\n", " ").strip()
        return re.sub(r"\s+", " ", text)

    def apply_data_masking(self, text: str) -> str:
        text = re.sub(r"[a-zA-Z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,}", "[EMAIL_MASKED]", text)
        text = re.sub(r"\b\d{9}\b", "[PHONE_MASKED]", text)
        return text

    # Carga de archivos con metadata
    def procesar_archivo(self, ruta_archivo, nivel_acceso="publico"):
        if not os.path.exists(ruta_archivo):
            logger.error(f"❌ El archivo no existe: {ruta_archivo}")
            return

        logger.info(f"📄 Procesando ({nivel_acceso}): {ruta_archivo}")
        documentos_langchain = []
        ext = ruta_archivo.lower()

        # Carga de archivos con metadata
        if ext.endswith(".pdf"):
            loader = PyPDFLoader(ruta_archivo)
            documentos_langchain = loader.load()
        elif ext.endswith(".docx"):
            doc_word = DocxDocument(ruta_archivo)
            texto_completo = "\n".join([para.text for para in doc_word.paragraphs])
            documentos_langchain = [Document(page_content=texto_completo, metadata={"source": ruta_archivo})]
        elif ext.endswith((".xlsx", ".xls")):
            df = pd.read_excel(ruta_archivo)
            texto_excel = df.to_string(index=False)
            documentos_langchain = [Document(page_content=texto_excel, metadata={"source": ruta_archivo})]
        elif ext.endswith(".txt"):
            with open(ruta_archivo, "r", encoding="utf-8") as f:
                texto_txt = f.read()
            documentos_langchain = [Document(page_content=texto_txt, metadata={"source": ruta_archivo})]
        else:
            logger.warning(f"⚠️ Formato {ext} no soportado.")
            return

        for doc in documentos_langchain:
            doc.page_content = self.apply_data_masking(self.clean_text(doc.page_content))
            # Forzamos que cada fragmento sepa a qué nivel pertenece
            doc.metadata["nivel_acceso"] = nivel_acceso 
            doc.metadata["source"] = os.path.basename(ruta_archivo)

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)
        trozos = text_splitter.split_documents(documentos_langchain)
        
        vectorstore = Qdrant(
            client=self.cliente_qdrant,
            collection_name=self.nombre_coleccion,
            embeddings=self.embeddings
        )
        vectorstore.add_documents(trozos)
        logger.info(f"✅ ¡{len(trozos)} fragmentos ({nivel_acceso}) indexados con éxito!")

if __name__ == "__main__":
    ingestor = IngestorDocumentos()
    
    # Ingesta organizada por carpetas o niveles
    archivos_publicos = ["manual_empleado.pdf", "normas_oficina.docx"]
    archivos_finanzas = ["auditoria_ventas.xlsx"]
    archivos_rrhh = ["salarios.txt"]
    
    for arc in archivos_publicos:
        ingestor.procesar_archivo(arc, nivel_acceso="publico")
        
    for arc in archivos_finanzas:
        ingestor.procesar_archivo(arc, nivel_acceso="finanzas")

    for arc in archivos_rrhh:
        ingestor.procesar_archivo(arc, nivel_acceso="rrhh")